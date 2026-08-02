# -*- coding: utf-8 -*-
"""Sinh BAO CAO TONG KET DE TAI NGHIEN CUU KHOA HOC hoan chinh.

Format hoc thuat: Times New Roman 13, gian dong 1.5, le 35/30/35/20mm; trang bia
ngoai + trong; Muc luc tu dong (TOC field), Danh muc hinh/bang/tu viet tat; Tom
tat; Mo dau; Chuong 1-3 viet day du; Ket luan; Tai lieu tham khao. Nhung hinh +
bang ty le %.

Sau khi luu, script tu mo Word (COM) de cap nhat muc luc va ghi lai (neu co Word).

Usage: python scripts/build_report_nckh.py
"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "outputs" / "report_figures"
CM = ROOT / "outputs" / "locbeef_rf_v1" / "confusion_matrix.png"
OUT = ROOT / "reports" / "bao_cao_nckh_nhan_biet_thit_tuoi.docx"

NAVY = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
style._element.rPr.rFonts.set(qn("w:cs"), FONT)
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)


def set_margins(section):
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(35)
    section.bottom_margin = Mm(30)
    section.left_margin = Mm(35)
    section.right_margin = Mm(20)


def _run(p, text, size=13, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), FONT)
    rf.set(qn("w:hAnsi"), FONT)
    rf.set(qn("w:cs"), FONT)
    rf.set(qn("w:eastAsia"), FONT)
    return r


def para(text="", align=None, bold=False, italic=False, size=13,
         indent=False, before=0, after=0, spacing=1.5, color=None):
    p = doc.add_paragraph()
    ppf = p.paragraph_format
    ppf.line_spacing = spacing
    ppf.space_before = Pt(before)
    ppf.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if indent:
        ppf.first_line_indent = Mm(10)
    if text:
        _run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def body(text):
    return para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True)


def h1(text, page_break=True, center=True):
    """Heading cap 1 (chuong / muc lon) — dung style Heading 1 de vao muc luc."""
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.5
    _run(p, text, size=14, bold=True, color=NAVY)
    return p


def h2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    _run(p, text, size=13, bold=True, color=NAVY)
    return p


def figure(path, caption, width_in=5.7):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(path), width=Inches(width_in))
    para(caption, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, before=2, after=8)


def grid_table(headers, rows, caption=None):
    if caption:
        para(caption, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, italic=True, size=12, before=6, after=2)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(c.paragraphs[0], h, bold=True, size=12)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cp = cells[i].paragraphs[0]
            if i == 0:
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(cp, v, size=12)
    para("", after=4)
    return t


def add_toc():
    p = doc.add_paragraph()
    r = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-2" \\h \\z \\u'
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    ft = OxmlElement("w:t"); ft.text = "Nhấn Ctrl+A rồi F9 để cập nhật mục lục."
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    r._r.append(fb); r._r.append(it); r._r.append(fs); r._r.append(ft); r._r.append(fe)


def page_num_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    r = p.add_run(); r.font.name = FONT; r.font.size = Pt(13)
    r._r.append(fb); r._r.append(it); r._r.append(fe)


# =====================================================================
# SECTION 1 - TRANG BIA (khong danh so trang)
# =====================================================================
sec1 = doc.sections[0]
set_margins(sec1)

para("[ TÊN TRƯỜNG / HỌC VIỆN ]", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, before=6)
para("[ KHOA / BỘ MÔN ]", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
para("---------------------------------------", WD_ALIGN_PARAGRAPH.CENTER)
para("[ Chèn logo tại đây ]", WD_ALIGN_PARAGRAPH.CENTER, italic=True, before=18, after=24)
para("BÁO CÁO TỔNG KẾT", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, before=6)
para("ĐỀ TÀI NGHIÊN CỨU KHOA HỌC", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, after=18)
para("NHẬN BIẾT THỊT TƯƠI BẰNG THỊ GIÁC MÁY TÍNH", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=17, after=4)
para("Phân loại độ tươi thịt bò từ ảnh bằng đặc trưng màu, texture và RandomForest",
     WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=13, after=24)
para("Sinh viên thực hiện: (điền họ tên)", WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("Giảng viên hướng dẫn: (điền học hàm, học vị, họ tên)", WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("Học phần: Thị giác máy tính", WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("HÀ NỘI - 2026", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, before=48)

# =====================================================================
# SECTION 2 - PHAN NOI DUNG (danh so trang tu 1)
# =====================================================================
sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
set_margins(sec2)
sec2.footer.is_linked_to_previous = False
sec1.footer.is_linked_to_previous = False
sec1.footer.paragraphs[0].text = ""
page_num_footer(sec2)
sectPr = sec2._sectPr
pgstart = OxmlElement("w:pgNumType"); pgstart.set(qn("w:start"), "1"); sectPr.append(pgstart)

# ---- MUC LUC ----
h1("MỤC LỤC", page_break=False)
add_toc()

# ---- DANH MUC TU VIET TAT ----
h1("DANH MỤC TỪ VIẾT TẮT")
grid_table(
    ["Từ viết tắt", "Nghĩa đầy đủ"],
    [
        ["RGB", "Red - Green - Blue (không gian màu)"],
        ["HSV", "Hue - Saturation - Value (không gian màu)"],
        ["CIELAB", "Không gian màu L*a*b* của CIE"],
        ["CLAHE", "Contrast Limited Adaptive Histogram Equalization"],
        ["LBP", "Local Binary Pattern (đặc trưng texture)"],
        ["RF", "Random Forest (rừng ngẫu nhiên)"],
        ["SVM", "Support Vector Machine (máy vector hỗ trợ)"],
        ["ML", "Machine Learning (học máy)"],
    ],
)

# ---- DANH MUC HINH VE ----
h1("DANH MỤC HÌNH VẼ")
for cap in [
    "Hình 2.1. Sơ đồ khối tổng quan của pipeline nhận biết thịt tươi.",
    "Hình 2.2. Quy trình tiền xử lý và tách vùng thịt.",
    "Hình 3.1. Ảnh mẫu từ bộ dữ liệu LocBeef.",
    "Hình 3.2. Phân bố màu vùng thịt trên mặt phẳng (a*, L*).",
    "Hình 3.3. Phân bố độ đỏ a* của thịt tươi và thịt hỏng.",
    "Hình 3.4. Confusion matrix trên tập test LocBeef.",
    "Hình 3.5. Đóng góp của từng nhóm đặc trưng vào mô hình.",
    "Hình 3.6. Giao diện web — kết quả cho ảnh thịt tươi.",
    "Hình 3.7. Giao diện web — kết quả cho ảnh thịt hỏng.",
    "Hình 3.8. Dự đoán của mô hình trên sáu ảnh test.",
]:
    para(cap, size=13, after=2)

# ---- DANH MUC BANG ----
h1("DANH MỤC BẢNG")
for cap in [
    "Bảng 1.1. So sánh hướng đặc trưng thủ công + ML và hướng học sâu.",
    "Bảng 2.1. Cấu trúc vector đặc trưng 174 chiều.",
    "Bảng 3.1. Phân bố bộ dữ liệu LocBeef.",
    "Bảng 3.2. Kết quả phân loại theo lớp trên tập test.",
    "Bảng 3.3. Độ quan trọng của các nhóm đặc trưng.",
]:
    para(cap, size=13, after=2)

# ---- TOM TAT ----
h1("TÓM TẮT")
body("Báo cáo trình bày một hệ thống thị giác máy tính nhận biết độ tươi của thịt bò từ ảnh RGB, phát biểu dưới dạng phân loại nhị phân: tươi (fresh) hoặc hỏng (spoiled). Hệ thống sử dụng cách tiếp cận học máy cổ điển với đặc trưng thủ công: sau khi cân bằng sáng bằng CLAHE và tách vùng thịt để loại nền, ảnh được biểu diễn bằng vector 174 chiều gồm histogram màu HSV, histogram Lab, các thống kê màu và đặc trưng texture LBP; bộ phân loại là RandomForest.")
body("Thực nghiệm trên bộ dữ liệu thật LocBeef gồm 3.268 ảnh thịt bò cho thấy mô hình đạt độ chính xác 97,9% trên tập kiểm tra 980 ảnh, với F1-score 97,8% cho lớp tươi và 97,9% cho lớp hỏng. Hệ thống được đóng gói thành ứng dụng web cho phép tải ảnh và nhận kết quả kèm xác suất. Báo cáo cũng phân tích trung thực các hạn chế về khả năng rò rỉ dữ liệu và lệch phân phối, cùng bài học từ việc loại bỏ một lớp hậu xử lý màu không hiệu quả sau khi đánh giá định lượng.")
para("Từ khóa: nhận biết thịt tươi, thị giác máy tính, đặc trưng màu, LBP, RandomForest, LocBeef.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=True, before=6)

# ---------------------------------------------------------- MO DAU
h1("MỞ ĐẦU")
h2("1. Lý do chọn đề tài")
body("Chất lượng và độ an toàn của thịt ảnh hưởng trực tiếp đến sức khỏe người tiêu dùng. Thịt là môi trường giàu dinh dưỡng nên rất dễ bị vi sinh vật phân hủy; quá trình ôi thiu kèm theo các biến đổi hóa học (oxy hóa myoglobin, sinh amin, mất nước) làm thay đổi màu sắc và kết cấu bề mặt. Bằng mắt thường, người mua thường đánh giá độ tươi qua màu đỏ/hồng, độ bóng và độ đàn hồi — đây chính là các tín hiệu thị giác mà máy tính có thể học để tự động hóa việc sàng lọc.")
body("So với kiểm nghiệm hóa sinh hoặc vi sinh, phương pháp dựa trên ảnh có ưu điểm không phá hủy mẫu, chi phí thấp, cho kết quả tức thì và dễ triển khai bằng camera phổ thông hoặc điện thoại. Bên cạnh ý nghĩa về sức khỏe, bài toán còn giúp giảm lãng phí thực phẩm, tối ưu luân chuyển hàng trong kho và giảm phụ thuộc vào đánh giá cảm quan chủ quan. Vì vậy, đề tài vừa có ý nghĩa khoa học, vừa có ý nghĩa thực tiễn và có tính khả thi khi tận dụng được các kỹ thuật xử lý ảnh, học máy sẵn có, chạy trên CPU mà không cần GPU.")
h2("2. Tổng quan tình hình nghiên cứu")
body("Đánh giá độ tươi thịt bằng ảnh đã được nghiên cứu theo hai hướng chính. Hướng đặc trưng thủ công kết hợp học máy sử dụng màu sắc (RGB/HSV/Lab), thống kê và texture (LBP, GLCM) làm đầu vào cho SVM, RandomForest hoặc k-NN; ưu điểm là nhẹ, dễ giải thích, phù hợp dữ liệu vừa và nhỏ. Hướng học sâu dùng mạng tích chập (ResNet, EfficientNet, MobileNet) hoặc Transformer để tự học đặc trưng, thường cho độ chính xác cao hơn khi có dữ liệu lớn và GPU.")
body("Điểm chung của các nghiên cứu là nhấn mạnh vai trò của màu sắc — đặc biệt độ đỏ liên quan tới myoglobin — như tín hiệu chủ đạo, đồng thời cảnh báo rủi ro mô hình học nhầm đặc trưng nền hoặc điều kiện chụp. Các công trình gần đây còn kết hợp phân vùng để tách vùng thịt và bổ sung cơ chế phát hiện mẫu ngoài phân phối. Qua khảo sát, có thể thấy các khoảng trống: (i) cần quy trình đặc trưng gọn nhưng đủ mạnh cho ảnh thịt thật; (ii) cần loại ảnh hưởng của nền; (iii) cần đánh giá định lượng trung thực trên dữ liệu thật thay vì heuristic cảm tính.")
h2("3. Mục tiêu nghiên cứu")
body("Mục tiêu tổng quát là xây dựng một hệ thống thị giác máy tính nhận biết độ tươi của thịt bò từ ảnh, phân loại nhị phân tươi/hỏng, kèm ứng dụng minh họa. Các mục tiêu cụ thể gồm: hệ thống hóa cơ sở lý thuyết liên quan; đề xuất quy trình tách vùng thịt và trích đặc trưng phù hợp; huấn luyện và đánh giá định lượng mô hình trên dữ liệu thật; đóng gói thành ứng dụng web sử dụng được.")
h2("4. Đối tượng và phạm vi nghiên cứu")
body("Đối tượng nghiên cứu: quy trình trích đặc trưng màu và texture cho ảnh thịt; kỹ thuật tách vùng thịt khỏi nền; các bộ phân loại học máy (RandomForest, SVM). Phạm vi: về nội dung, tập trung phân loại nhị phân bằng đặc trưng thủ công + học máy, không huấn luyện mô hình học sâu từ đầu; về dữ liệu, thực nghiệm trên bộ LocBeef gồm 3.268 ảnh thịt bò; về công cụ, hệ thống chạy trên CPU bằng Python và thư viện mã nguồn mở.")
h2("5. Phương pháp nghiên cứu")
body("Phương pháp nghiên cứu lý thuyết: thu thập, tổng hợp và phân tích tài liệu về không gian màu, CLAHE, LBP và các thuật toán phân loại. Phương pháp thực nghiệm: thiết kế và cài đặt pipeline; huấn luyện trên tập train và đánh giá trên tập test độc lập. Phương pháp đánh giá và thống kê: sử dụng accuracy, precision, recall, F1-score và confusion matrix; phân tích độ quan trọng đặc trưng và phân tích lỗi. Công cụ: Python, OpenCV, NumPy, scikit-learn, Matplotlib, Flask, Pillow, joblib.")
h2("6. Ý nghĩa khoa học và thực tiễn")
body("Về khoa học, đề tài đề xuất một quy trình đặc trưng gọn (174 chiều) kết hợp màu và texture trên vùng thịt đã tách nền, cùng đánh giá định lượng trung thực trên dữ liệu thật. Về thực tiễn, hệ thống cung cấp công cụ sàng lọc nhanh độ tươi thịt bằng ảnh, có ứng dụng web dễ sử dụng, chạy trên phần cứng phổ thông.")
h2("7. Bố cục báo cáo")
body("Ngoài phần mở đầu và kết luận, báo cáo gồm ba chương: Chương 1 trình bày cơ sở lý thuyết; Chương 2 trình bày phương pháp đề xuất; Chương 3 trình bày thực nghiệm, kết quả và ứng dụng web.")

# ---------------------------------------------------------- CHUONG 1
h1("CHƯƠNG 1. CƠ SỞ LÝ THUYẾT")
h2("1.1 Không gian màu")
body("Ảnh số thường được lưu trong không gian RGB, tuy nhiên RGB trộn lẫn thông tin độ sáng và sắc màu nên nhạy với thay đổi chiếu sáng. Không gian HSV tách riêng sắc độ (Hue), độ bão hòa (Saturation) và độ sáng (Value); sắc độ mô tả 'màu' của thịt gần như độc lập với độ sáng, phù hợp để phân biệt đỏ tươi với nâu/xám. Không gian CIELAB gồm L* (độ sáng), a* (trục đỏ–xanh lục) và b* (trục vàng–xanh lam); kênh a* đặc biệt quan trọng vì độ 'đỏ' của thịt liên hệ trực tiếp với trạng thái myoglobin — thịt tươi có a* cao, thịt hỏng có a* giảm.")
body("Chuyển RGB → HSV dựa trên giá trị lớn nhất và nhỏ nhất của ba kênh; chuyển RGB → CIELAB đi qua không gian trung gian XYZ rồi áp hàm phi tuyến xấp xỉ cảm nhận màu của mắt người, nhờ đó khoảng cách trong Lab gần với cảm nhận khác biệt màu hơn RGB. Lưu ý OpenCV biểu diễn kênh H trong khoảng 0–179 và các kênh còn lại trong 0–255.")
h2("1.2 Cân bằng sáng CLAHE")
body("CLAHE là kỹ thuật cân bằng histogram thích nghi theo từng vùng nhỏ của ảnh, có giới hạn tương phản để tránh khuếch đại nhiễu. Trong pipeline, CLAHE được áp trên kênh L của Lab rồi ghép lại, giúp giảm ảnh hưởng của chiếu sáng không đều (bóng, phản chiếu) mà vẫn giữ nguyên thông tin màu ở kênh a*, b*.")
h2("1.3 Đặc trưng texture LBP")
body("Local Binary Pattern (LBP) mô tả kết cấu cục bộ. Với mỗi điểm ảnh có giá trị trung tâm g_c và 8 lân cận g_0..g_7, mã LBP được tính bằng LBP = Σ (i=0..7) s(g_i − g_c)·2^i, trong đó s(x)=1 nếu x≥0 và s(x)=0 nếu ngược lại. Histogram của mã LBP trên toàn vùng thịt (32 bins) mô tả độ thô/mịn, vân cơ và các đốm bề mặt — những đặc điểm thay đổi khi thịt mất nước và bị oxy hóa. LBP bất biến với thay đổi độ sáng đơn điệu, bổ trợ tốt cho đặc trưng màu.")
h2("1.4 Các bộ phân loại")
body("RandomForest là tập hợp nhiều cây quyết định, mỗi cây huấn luyện trên một mẫu bootstrap và tại mỗi nút chỉ xét một tập con đặc trưng ngẫu nhiên; nút được tách sao cho giảm độ vẩn đục Gini = 1 − Σ_k p_k². Sự đa dạng giữa các cây làm giảm phương sai, giúp mô hình ổn định và ít overfit; xác suất dự đoán là trung bình tỉ lệ phiếu các cây. RandomForest mạnh với đặc trưng không đồng nhất, ít cần chuẩn hóa và cung cấp độ quan trọng đặc trưng. SVM kernel RBF tìm siêu phẳng lề cực đại trong không gian nâng chiều bởi hàm nhân Gaussian, nhưng cần chuẩn hóa và nhạy tham số; do đó đề tài chọn RandomForest làm mô hình chính.")
h2("1.5 Chỉ số đánh giá")
body("Gọi TP, TN, FP, FN là số dự đoán đúng-dương, đúng-âm, sai-dương và sai-âm. Khi đó Accuracy = (TP+TN)/(TP+TN+FP+FN); Precision = TP/(TP+FP); Recall = TP/(TP+FN); F1 = 2·Precision·Recall/(Precision+Recall). Confusion matrix cho biết lớp nào hay bị nhầm sang lớp nào. Với bài toán an toàn thực phẩm, recall của lớp 'hỏng' đặc biệt quan trọng vì bỏ sót thịt hỏng nguy hiểm hơn báo nhầm thịt tươi thành hỏng.")
h2("1.6 So sánh hai hướng tiếp cận")
body("Bảng 1.1 so sánh hướng đặc trưng thủ công + học máy với hướng học sâu, làm cơ sở cho lựa chọn của đề tài.")
grid_table(
    ["Tiêu chí", "Đặc trưng thủ công + ML", "Học sâu (CNN/Transformer)"],
    [
        ["Nhu cầu dữ liệu", "Vừa và nhỏ", "Lớn"],
        ["Tài nguyên", "CPU là đủ", "Thường cần GPU"],
        ["Khả năng giải thích", "Cao (đặc trưng rõ ràng)", "Thấp (hộp đen)"],
        ["Độ chính xác tiềm năng", "Khá – cao", "Cao khi đủ dữ liệu"],
        ["Thời gian triển khai", "Nhanh", "Lâu hơn"],
    ],
    caption="Bảng 1.1. So sánh hướng đặc trưng thủ công + ML và hướng học sâu.",
)
body("Với phạm vi một đề tài nghiên cứu khoa học sinh viên, hướng đặc trưng thủ công + RandomForest là lựa chọn hợp lý: nhẹ, dễ giải thích và vẫn đạt độ chính xác cao trên bộ dữ liệu đang xét.")
para("Kết luận Chương 1: chương đã hệ thống hóa các khái niệm nền tảng về không gian màu, cân bằng sáng, đặc trưng texture và bộ phân loại, làm cơ sở cho phương pháp đề xuất ở Chương 2.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=True, before=6)

# ---------------------------------------------------------- CHUONG 2
h1("CHƯƠNG 2. PHƯƠNG PHÁP ĐỀ XUẤT")
h2("2.1 Tổng quan pipeline")
body("Hệ thống được thiết kế thành sáu bước xử lý tuần tự: đọc ảnh và resize về 224×224; cân bằng sáng bằng CLAHE; tách vùng thịt; trích đặc trưng màu và texture trên vùng thịt; ghép thành vector 174 chiều; phân loại bằng RandomForest và đánh giá. Sơ đồ khối tổng quan được trình bày ở Hình 2.1.")
figure(FIG / "pipeline_diagram.png", "Hình 2.1. Sơ đồ khối tổng quan của pipeline nhận biết thịt tươi.", 6.4)
h2("2.2 Tiền xử lý")
body("Ảnh đầu vào được resize về kích thước cố định 224×224 để chuẩn hóa. Sau đó, CLAHE được áp trên kênh L của Lab với clipLimit = 2,0 và ô lưới 8×8 nhằm giảm ảnh hưởng của chiếu sáng không đều mà vẫn giữ thông tin màu ở kênh a*, b*.")
h2("2.3 Tách vùng thịt")
body("Bước tách vùng thịt tạo một mask nhị phân giữ lại pixel thuộc miếng thịt. Quy tắc: loại các pixel quá sáng và ít bão hòa (S < 30 và V > 220 — tương ứng đĩa/nền trắng) và pixel quá tối (V ≤ 20 — vùng bóng). Mask được làm sạch bằng phép hình thái học mở rồi đóng với phần tử cấu trúc elip 9×9: phép mở loại đốm nhiễu nhỏ, phép đóng lấp các lỗ bên trong miếng thịt. Nếu vùng thịt phát hiện được nhỏ hơn 3% diện tích ảnh, hệ thống bỏ mask và dùng toàn ảnh để tránh mất dữ liệu. Toàn bộ đặc trưng màu chỉ tính trên vùng thịt, nhờ đó mô hình không học nhầm màu nền hay màu đĩa (Hình 2.2).")
figure(FIG / "preprocessing.png", "Hình 2.2. Quy trình tiền xử lý và tách vùng thịt.", 6.5)
h2("2.4 Trích đặc trưng")
body("Vector đặc trưng gồm bốn nhóm, tổng cộng 174 chiều (Bảng 2.1). Histogram từng kênh được chuẩn hóa theo mật độ để bất biến với số lượng pixel vùng thịt. Đặc trưng màu (HSV, Lab, thống kê) mô tả sự khác biệt về sắc độ và độ đỏ giữa thịt tươi và hỏng; đặc trưng LBP mô tả vân và độ thô mịn bề mặt.")
grid_table(
    ["Nhóm đặc trưng", "Chi tiết", "Số chiều", "Tỷ lệ"],
    [
        ["Histogram HSV", "H (32), S (16), V (16)", "64", "36,8%"],
        ["Histogram Lab", "L (16), a (16), b (16)", "48", "27,6%"],
        ["Thống kê màu", "mean, std, p10/p50/p90 × 6 kênh", "30", "17,2%"],
        ["Texture LBP", "LBP 32 bins", "32", "18,4%"],
        ["Tổng cộng", "", "174", "100%"],
    ],
    caption="Bảng 2.1. Cấu trúc vector đặc trưng 174 chiều.",
)
h2("2.5 Mô hình phân loại")
body("Mô hình chính là RandomForest với 300 cây, class_weight='balanced' và random_state cố định để tái lập. Mô hình được huấn luyện trên tập train và đánh giá trên tập test độc lập; riêng mô hình đóng gói kèm ứng dụng được huấn luyện lại trên toàn bộ dữ liệu để tận dụng tối đa mẫu. Đây là thực hành phổ biến: dùng phần held-out để ước lượng khả năng tổng quát, rồi huấn luyện mô hình triển khai trên toàn bộ dữ liệu.")
para("Kết luận Chương 2: chương đã trình bày pipeline sáu bước với đóng góp chính là bước tách vùng thịt và tập đặc trưng 174 chiều kết hợp màu và texture, cùng lựa chọn mô hình RandomForest.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=True, before=6)

# ---------------------------------------------------------- CHUONG 3
h1("CHƯƠNG 3. THỰC NGHIỆM VÀ KẾT QUẢ")
h2("3.1 Bộ dữ liệu")
body("Thực nghiệm sử dụng bộ LocBeef — Beef Quality Image Dataset (Kaggle), gồm 3.268 ảnh thịt bò địa phương vùng Aceh, hai lớp fresh và rotten (ánh xạ sang spoiled), đã chia sẵn train/test. Ảnh chụp miếng thịt trên đĩa/nền gỗ trong điều kiện ánh sáng tương đối đồng nhất. Hai lớp cân bằng hoàn hảo (mỗi lớp 1.634 ảnh, chiếm 50,0%), nên accuracy là chỉ số đánh giá hợp lý (Bảng 3.1).")
grid_table(
    ["Tập", "fresh", "rotten", "Tổng", "Tỷ lệ"],
    [
        ["Train", "1.144", "1.144", "2.288", "70,0%"],
        ["Test", "490", "490", "980", "30,0%"],
        ["Tổng", "1.634", "1.634", "3.268", "100%"],
    ],
    caption="Bảng 3.1. Phân bố bộ dữ liệu LocBeef.",
)
body("Do kho ảnh gốc lớn (khoảng 5,5 GB), quá trình huấn luyện đọc ảnh trực tiếp từ file nén trong bộ nhớ mà không giải nén ra ổ đĩa. Một số ảnh mẫu và phân bố màu của hai lớp được minh họa ở Hình 3.1, Hình 3.2 và Hình 3.3.")
figure(FIG / "dataset_samples.png", "Hình 3.1. Ảnh mẫu từ bộ dữ liệu LocBeef (trên: tươi, dưới: hỏng).", 6.2)
figure(FIG / "eda_scatter.png", "Hình 3.2. Phân bố màu vùng thịt trên mặt phẳng (a*, L*).", 5.3)
figure(FIG / "color_distribution.png", "Hình 3.3. Phân bố độ đỏ a* của thịt tươi và thịt hỏng.", 5.5)
h2("3.2 Thiết lập thực nghiệm")
body("Mô hình được huấn luyện trên 2.288 ảnh train và đánh giá trên 980 ảnh test độc lập, qua đúng pipeline mô tả ở Chương 2. Các chỉ số accuracy, precision, recall, F1-score và confusion matrix được tính trên tập test.")
h2("3.3 Kết quả tổng thể")
body("Trên tập test 980 ảnh, mô hình đạt accuracy 97,9% (959/980 ảnh đúng, 21 lỗi). Kết quả theo lớp được trình bày dưới dạng phần trăm ở Bảng 3.2.")
grid_table(
    ["Chỉ số", "fresh", "spoiled"],
    [
        ["Precision", "100,0%", "95,9%"],
        ["Recall", "95,7%", "100,0%"],
        ["F1-score", "97,8%", "97,9%"],
        ["Accuracy chung", "97,9%", "97,9%"],
    ],
    caption="Bảng 3.2. Kết quả phân loại theo lớp trên tập test.",
)
h2("3.4 Ma trận nhầm lẫn")
body("Confusion matrix (Hình 3.4) cho thấy toàn bộ 490 ảnh hỏng được nhận đúng và 469 ảnh tươi được nhận đúng; 21 ảnh tươi bị đoán nhầm thành hỏng, không có ảnh hỏng nào bị đoán nhầm thành tươi. Mô hình nghiêng về phía thận trọng (thiên báo hỏng) — khi sai, nó sai theo hướng cảnh báo hỏng chứ không bỏ sót thịt hỏng, đặc tính có lợi cho an toàn thực phẩm.")
figure(CM, "Hình 3.4. Confusion matrix trên tập test LocBeef (980 ảnh).", 4.5)
h2("3.5 Độ quan trọng đặc trưng")
body("Tổng hợp độ quan trọng của RandomForest theo nhóm (Bảng 3.3, Hình 3.5) cho thấy các đặc trưng màu chiếm tới 96,4% quyết định, trong khi texture LBP chỉ đóng góp 3,6%. Kết quả phù hợp trực giác: màu là dấu hiệu chính phân biệt thịt tươi và hỏng.")
grid_table(
    ["Nhóm đặc trưng", "Độ quan trọng"],
    [
        ["Histogram HSV", "56,5%"],
        ["Thống kê màu", "27,4%"],
        ["Histogram Lab", "12,6%"],
        ["Texture LBP", "3,6%"],
        ["Tổng nhóm màu", "96,4%"],
    ],
    caption="Bảng 3.3. Độ quan trọng của các nhóm đặc trưng.",
)
figure(FIG / "feature_importance.png", "Hình 3.5. Đóng góp của từng nhóm đặc trưng vào mô hình.", 5.6)
h2("3.6 Phân tích lỗi")
body("Cả 21 lỗi đều là ảnh tươi bị phân loại thành hỏng, thường do miếng thịt sẫm màu hơn trung bình hoặc bị bóng/thiếu sáng cục bộ khiến độ đỏ đo được thấp đi — các trường hợp ở vùng chuyển tiếp về màu. Hướng khắc phục gồm tăng cường dữ liệu và bổ sung đặc trưng phản ánh độ đồng nhất màu trong vùng thịt.")
h2("3.7 Ứng dụng web demo")
body("Hệ thống được đóng gói thành ứng dụng web bằng Flask: người dùng tải ảnh lên trình duyệt và nhận nhãn tươi/hỏng kèm xác suất từng lớp. Ứng dụng có cơ chế đọc cả định dạng ảnh AVIF/HEIC (thường gặp trên điện thoại) bằng Pillow khi OpenCV không giải mã được. Hình 3.6 và Hình 3.7 minh họa giao diện thực tế khi phân tích ảnh thịt tươi và thịt hỏng; Hình 3.8 tổng hợp dự đoán trên sáu ảnh test.")
figure(FIG / "web_result_fresh.png", "Hình 3.6. Giao diện web — kết quả cho ảnh thịt tươi.", 6.0)
figure(FIG / "web_result_spoiled.png", "Hình 3.7. Giao diện web — kết quả cho ảnh thịt hỏng.", 6.0)
figure(FIG / "demo_predictions.png", "Hình 3.8. Dự đoán của mô hình trên sáu ảnh test (đều đúng).", 6.2)
h2("3.8 Đánh giá khả năng tổng quát hóa")
body("Để đánh giá khả năng tổng quát ngoài miền huấn luyện, mô hình được thử với một số ảnh thịt lấy ngẫu nhiên trên web (khác camera, ánh sáng, cách trình bày). Kết quả cho thấy độ tin cậy giảm rõ và xuất hiện lỗi — ví dụ một ảnh thịt đã ngả màu bị phân loại là tươi. Điều này khẳng định mô hình mạnh trên ảnh giống phân phối huấn luyện nhưng chưa tổng quát cho mọi điều kiện.")
para("Kết luận Chương 3: mô hình đạt 97,9% accuracy trên tập test LocBeef với đặc trưng màu đóng vai trò quyết định; hệ thống hoạt động tốt qua ứng dụng web, đồng thời bộc lộ giới hạn về khả năng tổng quát cần được lưu ý.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=True, before=6)

# ---------------------------------------------------------- KET LUAN
h1("KẾT LUẬN VÀ KIẾN NGHỊ")
h2("1. Kết quả đạt được")
body("Đề tài đã xây dựng hoàn chỉnh một hệ thống thị giác máy tính nhận biết độ tươi thịt bò từ ảnh: (i) hệ thống hóa cơ sở lý thuyết về đặc trưng màu, CLAHE, LBP và bộ phân loại học máy; (ii) đề xuất pipeline tách vùng thịt + đặc trưng 174 chiều + RandomForest; (iii) đạt 97,9% accuracy trên tập test thật LocBeef và đóng gói thành ứng dụng web sử dụng được.")
h2("2. Hạn chế")
body("Kết quả dựa trên phân chia train/test có sẵn của bộ dữ liệu nên có thể lạc quan nếu tồn tại ảnh cùng mẫu vật ở cả hai tập. Mô hình bị giảm hiệu năng khi gặp ảnh lệch phân phối (khác camera, ánh sáng, loại thịt). Ngoài ra, một lớp hậu xử lý màu 'hybrid' từng được thử nghiệm nhưng bị loại bỏ do làm accuracy tụt xuống 50% trên dữ liệu thật — minh chứng cho tầm quan trọng của việc kiểm chứng heuristic bằng đánh giá định lượng.")
h2("3. Hướng phát triển")
body("Các hướng phát triển tiếp theo gồm: chia dữ liệu theo từng mẫu vật để loại trừ rò rỉ và ước lượng tổng quát chính xác hơn; tăng cường và đa dạng hóa dữ liệu; áp dụng transfer learning (MobileNetV3, EfficientNet-B0) và so sánh với baseline hiện tại; bổ sung cơ chế từ chối dự đoán khi ảnh ngoài phân phối hoặc độ tin cậy thấp. Hệ thống là công cụ hỗ trợ sàng lọc, không thay thế kiểm nghiệm vi sinh chính thức.")

# ---------------------------------------------------------- TAI LIEU
h1("TÀI LIỆU THAM KHẢO")


def ref(text):
    p = doc.add_paragraph()
    ppf = p.paragraph_format
    ppf.line_spacing = 1.5
    ppf.space_after = Pt(2)
    ppf.left_indent = Mm(10)
    ppf.first_line_indent = Mm(-10)
    _run(p, text, size=13)
    return p


para("Tiếng Anh", bold=True, before=4)
ref("[1] Ojala T., Pietikäinen M., Mäenpää T. (2002), “Multiresolution gray-scale and rotation invariant texture classification with local binary patterns”, IEEE TPAMI, 24(7), pp. 971-987.")
ref("[2] Breiman L. (2001), “Random Forests”, Machine Learning, 45(1), pp. 5-32.")
ref("[3] Zuiderveld K. (1994), “Contrast Limited Adaptive Histogram Equalization”, Graphics Gems IV, pp. 474-485.")
ref("[4] Bramantyo H. A., et al. (2026), “Deep Learning-Based Meat Freshness Detection with Segmentation and OOD-Aware Classification”, arXiv:2603.00368.")
ref("[5] Hidalgo M. M., et al. (2025), “Leveraging pre-trained computer vision models for accurate classification of meat freshness”, Food Chemistry, 495(Pt 3), 146430.")
para("Tài liệu Internet", bold=True, before=6)
ref("[6] LocBeef — Beef Quality Image Dataset. Kaggle. <https://www.kaggle.com/datasets/mexwell/locbeef-beef-quality-image-dataset>. Truy cập ngày 02 tháng 8 năm 2026.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print("Saved:", OUT)

# ---- Cap nhat muc luc bang Word COM (neu co) ----
try:
    import win32com.client as win32
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    d = word.Documents.Open(str(OUT))
    d.Fields.Update()
    if d.TablesOfContents.Count > 0:
        d.TablesOfContents(1).Update()
    d.Repaginate()
    pages = d.ComputeStatistics(2)
    d.Save()
    d.Close(False)
    word.Quit()
    print("TOC updated. Pages:", pages)
except Exception as e:  # noqa: BLE001
    print("(Bo qua cap nhat TOC - khong co Word):", e)
