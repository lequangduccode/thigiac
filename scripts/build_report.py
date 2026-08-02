from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = ROOT / "reports" / "bao_cao_nhan_biet_thit_tuoi.docx"
FIG = ROOT / "outputs" / "report_figures"
CONFUSION_MATRIX = ROOT / "outputs" / "locbeef_rf_v1" / "confusion_matrix.png"


def set_font(run, name: str = "Calibri", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    for run in paragraph.runs:
        set_font(run, size=16 if level == 1 else 13, bold=True)
        run.font.color.rgb = RGBColor(46, 116, 181) if level < 3 else RGBColor(31, 77, 120)
    return paragraph


def add_paragraph(doc: Document, text: str = "", bold_prefix: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, size=11, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_font(run, size=11)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=11)
    return paragraph


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.4
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, size=11, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_font(run, size=11)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=11)
    return paragraph


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.4
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_font(run, size=11)
    return paragraph


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "1F4E79")
        set_cell_width(cell, widths[index])
        run = cell.paragraphs[0].add_run(header)
        set_font(run, size=10, bold=True)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_width(cells[index], widths[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_i % 2 == 1:
                set_cell_shading(cells[index], "EEF3FA")
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_font(run, size=10)
    doc.add_paragraph()
    return table


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.6) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cap_run = cap.add_run(caption)
    set_font(cap_run, size=9)
    cap_run.italic = True
    cap_run.font.color.rgb = RGBColor(90, 90, 90)


def build() -> None:
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    # ---------------------------------------------------------------- COVER
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("NHẬN BIẾT THỊT TƯƠI\nBẰNG THỊ GIÁC MÁY TÍNH")
    set_font(run, size=26, bold=True)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Phân loại độ tươi thịt bò từ ảnh bằng đặc trưng màu + texture,\ntách vùng thịt và RandomForest")
    set_font(run, size=13)
    run.font.color.rgb = RGBColor(85, 85, 85)

    for _ in range(6):
        doc.add_paragraph()
    for line in [
        "Báo cáo bài tập môn: Thị giác máy tính",
        "Bộ dữ liệu: LocBeef — Beef Quality Image Dataset (3.268 ảnh)",
        "Mô hình: RandomForest — Accuracy 97.9% trên test set",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(line)
        set_font(r, size=12, bold=True)
        r.font.color.rgb = RGBColor(46, 116, 181)

    doc.add_section(WD_SECTION.NEW_PAGE)

    # -------------------------------------------------------------- ABSTRACT
    add_heading(doc, "Tóm tắt", 1)
    add_paragraph(
        doc,
        "Báo cáo trình bày một hệ thống thị giác máy tính nhận biết độ tươi của thịt bò từ ảnh RGB. "
        "Bài toán được phát biểu dưới dạng phân loại nhị phân: tươi (fresh) hoặc hỏng (spoiled). "
        "Hệ thống dùng cách tiếp cận học máy cổ điển với đặc trưng thủ công: sau khi cân bằng sáng bằng CLAHE và "
        "tách vùng thịt để loại nền, ảnh được biểu diễn bằng vector 174 chiều gồm histogram màu HSV/Lab, các thống kê "
        "màu và đặc trưng texture LBP; bộ phân loại là RandomForest. Trên bộ dữ liệu thật LocBeef gồm 3.268 ảnh thịt "
        "bò địa phương, mô hình đạt độ chính xác 97.9% trên tập kiểm tra 980 ảnh. Hệ thống được đóng gói thành một "
        "ứng dụng web cho phép người dùng tải ảnh và nhận kết quả dự đoán kèm xác suất. Báo cáo cũng nêu trung thực "
        "các hạn chế quan trọng — khả năng rò rỉ dữ liệu do cách chia sẵn của bộ dữ liệu và hiện tượng lệch phân phối "
        "khi áp dụng cho ảnh ngoài miền huấn luyện — cùng bài học rút ra từ việc loại bỏ một lớp hậu xử lý màu không "
        "hiệu quả sau khi đánh giá định lượng.",
    )

    add_heading(doc, "Mục lục", 1)
    for i, name in enumerate([
        "Giới thiệu", "Cơ sở lý thuyết", "Công trình liên quan", "Bộ dữ liệu",
        "Phương pháp đề xuất", "Thực nghiệm và kết quả", "Ứng dụng web demo",
        "Chi tiết cài đặt và khả năng tái lập", "Thảo luận",
        "Kết luận và hướng phát triển",
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{i}.  {name}")
        set_font(r, size=11)

    doc.add_section(WD_SECTION.NEW_PAGE)

    # ---------------------------------------------------------- 1. GIOI THIEU
    add_heading(doc, "1. Giới thiệu", 1)
    add_heading(doc, "1.1 Bối cảnh", 2)
    add_paragraph(
        doc,
        "Chất lượng và độ an toàn của thịt là mối quan tâm trực tiếp đối với sức khỏe người tiêu dùng. Thịt là môi "
        "trường giàu dinh dưỡng nên rất dễ bị vi sinh vật phân hủy; quá trình ôi thiu kèm theo các thay đổi hóa học "
        "(oxy hóa myoglobin, sinh amin, mất nước) làm biến đổi màu sắc và kết cấu bề mặt. Bằng mắt thường, người mua "
        "thường đánh giá độ tươi qua màu đỏ/hồng tươi, độ bóng và độ đàn hồi. Đây chính là những tín hiệu thị giác mà "
        "một hệ thống thị giác máy tính có thể học để tự động hóa việc sàng lọc nhanh.",
    )
    add_paragraph(
        doc,
        "So với các phương pháp kiểm nghiệm hóa sinh hay vi sinh, phương pháp dựa trên ảnh có ưu điểm không phá hủy "
        "mẫu, chi phí thấp, cho kết quả tức thì và dễ triển khai bằng camera phổ thông hoặc điện thoại. Vì vậy nó phù "
        "hợp cho các ứng dụng hỗ trợ tại siêu thị, bếp ăn công nghiệp, chợ đầu mối hay kho lạnh.",
    )
    add_paragraph(
        doc,
        "Bên cạnh khía cạnh sức khỏe, bài toán còn có ý nghĩa kinh tế và môi trường: sàng lọc sớm giúp giảm lãng phí "
        "thực phẩm, tối ưu luân chuyển hàng trong kho và hỗ trợ minh bạch chất lượng cho người mua. Việc tự động hóa "
        "bằng camera cũng giảm phụ thuộc vào đánh giá cảm quan chủ quan của từng người, vốn dễ thiếu nhất quán.",
    )
    add_heading(doc, "1.2 Phát biểu bài toán", 2)
    add_paragraph(
        doc,
        "Đầu vào là một ảnh màu RGB chụp miếng thịt bò trong điều kiện ánh sáng thông thường. Đầu ra là nhãn phân loại "
        "thuộc một trong hai lớp: fresh (tươi) hoặc spoiled (hỏng/ôi). Đây là bài toán phân loại nhị phân, được lựa "
        "chọn để khớp với nhãn của bộ dữ liệu thật sử dụng trong báo cáo (LocBeef gồm hai lớp fresh và rotten). Ngoài "
        "nhãn, hệ thống còn trả về xác suất thuộc từng lớp để người dùng ước lượng mức độ tin cậy.",
    )
    add_heading(doc, "1.3 Mục tiêu", 2)
    add_bullet(doc, "Xây dựng pipeline thị giác máy tính hoàn chỉnh: tiền xử lý, tách vùng thịt, trích đặc trưng, phân loại và đánh giá.")
    add_bullet(doc, "Huấn luyện và đánh giá mô hình trên bộ dữ liệu ảnh thịt thật thay vì dữ liệu mô phỏng.")
    add_bullet(doc, "Đóng gói thành ứng dụng web demo cho phép tải ảnh và xem kết quả trực quan.")
    add_bullet(doc, "Phân tích trung thực điểm mạnh, hạn chế và rủi ro của hệ thống.")
    add_heading(doc, "1.4 Đóng góp", 2)
    add_bullet(doc, "Một pipeline đặc trưng thủ công 174 chiều kết hợp màu (HSV, Lab), thống kê màu và texture (LBP), tính trên vùng thịt sau khi loại nền.")
    add_bullet(doc, "Kết quả thực nghiệm 97.9% accuracy trên tập kiểm tra LocBeef, kèm phân tích theo lớp, độ quan trọng đặc trưng và phân tích lỗi.")
    add_bullet(doc, "Ứng dụng web Flask có khả năng đọc cả định dạng ảnh AVIF/HEIC từ điện thoại.")
    add_bullet(doc, "Một bài học kỹ thuật: minh chứng vì sao heuristic cảm tính cần được kiểm chứng bằng đánh giá định lượng trên dữ liệu thật.")
    add_heading(doc, "1.5 Cấu trúc báo cáo", 2)
    add_paragraph(
        doc,
        "Phần 2 trình bày cơ sở lý thuyết về không gian màu, CLAHE, LBP và các bộ phân loại. Phần 3 điểm qua công "
        "trình liên quan. Phần 4 mô tả bộ dữ liệu. Phần 5 trình bày phương pháp đề xuất. Phần 6 báo cáo thực nghiệm "
        "và kết quả. Phần 7 giới thiệu ứng dụng web demo. Phần 8 thảo luận hạn chế và phần 9 kết luận.",
    )

    # ------------------------------------------------------- 2. CO SO LY THUYET
    add_heading(doc, "2. Cơ sở lý thuyết", 1)
    add_heading(doc, "2.1 Không gian màu", 2)
    add_paragraph(
        doc,
        "Ảnh số thường được lưu trong không gian RGB. Tuy nhiên RGB trộn lẫn thông tin về độ sáng và sắc màu nên nhạy "
        "với thay đổi chiếu sáng. Trong bài toán đánh giá màu thịt, hai không gian màu hữu ích hơn được sử dụng:",
    )
    add_bullet(doc, "HSV (Hue, Saturation, Value): tách riêng sắc độ (H), độ bão hòa (S) và độ sáng (V). Sắc độ H mô tả 'màu' của thịt gần như độc lập với độ sáng, rất phù hợp để phân biệt đỏ tươi với nâu/xám.", bold_prefix="HSV (Hue, Saturation, Value):")
    add_bullet(doc, "CIELAB (L*, a*, b*): L* là độ sáng, a* là trục đỏ–xanh lục, b* là trục vàng–xanh lam. Kênh a* đặc biệt quan trọng vì độ 'đỏ' của thịt liên hệ trực tiếp với trạng thái của myoglobin; thịt tươi có a* cao, thịt hỏng có a* giảm.", bold_prefix="CIELAB (L*, a*, b*):")
    add_paragraph(
        doc,
        "OpenCV biểu diễn kênh H trong khoảng 0–179 (thay vì 0–359 độ) và các kênh còn lại trong 0–255; báo cáo lưu ý "
        "điều này khi thiết lập số bins cho histogram.",
    )
    add_paragraph(
        doc,
        "Chuyển đổi RGB → HSV dựa trên giá trị lớn nhất (V = max(R,G,B)), độ bão hòa S = (V − min)/V và sắc độ H tính "
        "từ kênh trội. Chuyển RGB → CIELAB đi qua không gian trung gian XYZ rồi áp hàm phi tuyến để xấp xỉ cảm nhận màu "
        "của mắt người; nhờ đó khoảng cách trong Lab gần với cảm nhận khác biệt màu hơn RGB. Tính chất 'đều theo cảm "
        "nhận' này là lý do Lab (đặc biệt kênh a*) rất phù hợp để lượng hóa độ đỏ của thịt.",
    )
    add_heading(doc, "2.2 Cân bằng sáng CLAHE", 2)
    add_paragraph(
        doc,
        "CLAHE (Contrast Limited Adaptive Histogram Equalization) là kỹ thuật cân bằng sáng thích nghi theo từng vùng "
        "nhỏ của ảnh, có giới hạn tương phản để tránh khuếch đại nhiễu. Trong pipeline, CLAHE được áp trên kênh L của "
        "không gian Lab rồi ghép lại, giúp giảm ảnh hưởng của chiếu sáng không đều (bóng, phản chiếu) mà vẫn giữ nguyên "
        "thông tin màu ở kênh a*, b*.",
    )
    add_heading(doc, "2.3 Histogram màu", 2)
    add_paragraph(
        doc,
        "Histogram màu đếm tần suất giá trị điểm ảnh trên mỗi kênh, tạo thành mô tả thống kê về phân bố màu của vùng "
        "quan tâm. Histogram bất biến với vị trí và hình dạng của miếng thịt trong khung hình, nên phù hợp khi tư thế "
        "và bố cục ảnh thay đổi. Báo cáo dùng histogram 1 chiều cho từng kênh (chuẩn hóa theo mật độ) để vector đặc "
        "trưng gọn và ổn định.",
    )
    add_heading(doc, "2.4 Đặc trưng texture LBP", 2)
    add_paragraph(
        doc,
        "Local Binary Pattern (LBP) mô tả kết cấu cục bộ. Với mỗi điểm ảnh, LBP so sánh 8 điểm lân cận với điểm trung "
        "tâm: lân cận sáng hơn cho bit 1, tối hơn cho bit 0, ghép 8 bit thành một mã 0–255. Histogram của các mã LBP "
        "trên toàn vùng thịt mô tả độ thô/mịn, vân cơ và các đốm bề mặt — những đặc điểm thay đổi khi thịt mất nước và "
        "bị oxy hóa. LBP có ưu điểm bất biến với thay đổi độ sáng đơn điệu, bổ trợ tốt cho đặc trưng màu.",
    )
    add_paragraph(
        doc,
        "Về mặt công thức, với điểm trung tâm giá trị g_c và 8 lân cận g_0..g_7, mã LBP được tính bằng "
        "LBP = Σ (i=0..7) s(g_i − g_c) · 2^i, trong đó s(x) = 1 nếu x ≥ 0 và s(x) = 0 nếu ngược lại. Histogram của mã "
        "LBP trên toàn vùng thịt (32 bins) tạo thành đặc trưng texture cuối cùng.",
    )
    add_heading(doc, "2.5 Bộ phân loại", 2)
    add_paragraph(
        doc,
        "RandomForest là tập hợp nhiều cây quyết định huấn luyện trên các mẫu bootstrap và tập con đặc trưng ngẫu nhiên; "
        "dự đoán bằng biểu quyết đa số. Ưu điểm là mạnh với đặc trưng không đồng nhất, ít cần chuẩn hóa, khó overfit khi "
        "đủ cây và cung cấp được độ quan trọng đặc trưng. Mã nguồn cũng hỗ trợ SVM kernel RBF (kèm chuẩn hóa) như một "
        "lựa chọn thay thế. Tham số class_weight='balanced' giúp cân bằng khi hai lớp lệch nhau (dù LocBeef vốn cân bằng).",
    )
    add_paragraph(
        doc,
        "Mỗi cây trong rừng được huấn luyện trên một mẫu bootstrap (lấy có hoàn lại) của tập train, và tại mỗi nút chỉ "
        "xét một tập con đặc trưng ngẫu nhiên (thường cỡ căn bậc hai của số chiều). Nút được tách sao cho giảm độ vẩn "
        "đục Gini, với Gini = 1 − Σ_k p_k², trong đó p_k là tỉ lệ mẫu thuộc lớp k tại nút. Sự đa dạng giữa các cây (do "
        "bootstrap + đặc trưng ngẫu nhiên) làm giảm phương sai của mô hình tổng hợp, giúp RandomForest ổn định và ít "
        "overfit. Xác suất dự đoán được ước lượng bằng trung bình tỉ lệ phiếu của các cây.",
    )
    add_paragraph(
        doc,
        "SVM kernel RBF, ngược lại, tìm siêu phẳng phân tách với lề cực đại trong không gian đặc trưng được nâng chiều "
        "ngầm bởi hàm nhân Gaussian K(x, x') = exp(−γ‖x − x'‖²). SVM thường cần chuẩn hóa đặc trưng (StandardScaler) và "
        "nhạy với tham số C, γ; đó là lý do baseline chính chọn RandomForest cho tính tiện dụng và độ bền.",
    )
    add_heading(doc, "2.6 Chỉ số đánh giá", 2)
    add_bullet(doc, "Accuracy: tỉ lệ dự đoán đúng trên toàn bộ tập kiểm tra.", bold_prefix="Accuracy:")
    add_bullet(doc, "Precision: trong các ảnh bị gán một lớp, bao nhiêu phần trăm thực sự thuộc lớp đó.", bold_prefix="Precision:")
    add_bullet(doc, "Recall: trong các ảnh thực sự thuộc một lớp, bao nhiêu phần trăm được nhận đúng.", bold_prefix="Recall:")
    add_bullet(doc, "F1-score: trung bình điều hòa của precision và recall.", bold_prefix="F1-score:")
    add_bullet(doc, "Confusion matrix: ma trận nhầm lẫn cho biết lớp nào hay bị nhầm sang lớp nào.", bold_prefix="Confusion matrix:")
    add_paragraph(
        doc,
        "Gọi TP, TN, FP, FN lần lượt là số dự đoán đúng-dương, đúng-âm, sai-dương và sai-âm. Khi đó: "
        "Accuracy = (TP + TN) / (TP + TN + FP + FN); Precision = TP / (TP + FP); Recall = TP / (TP + FN); "
        "F1 = 2 · Precision · Recall / (Precision + Recall). Với bài toán an toàn thực phẩm, recall của lớp 'hỏng' "
        "đặc biệt quan trọng vì bỏ sót thịt hỏng (FN) nguy hiểm hơn báo nhầm thịt tươi thành hỏng (FP).",
    )

    # ---------------------------------------------------- 3. CONG TRINH LIEN QUAN
    add_heading(doc, "3. Công trình liên quan", 1)
    add_paragraph(
        doc,
        "Đánh giá độ tươi thịt bằng ảnh đã được nghiên cứu theo hai hướng chính. Hướng đặc trưng thủ công + học máy sử "
        "dụng màu sắc (RGB/HSV/Lab), thống kê và texture (LBP, GLCM) làm đầu vào cho SVM, RandomForest hoặc k-NN; ưu "
        "điểm là nhẹ, dễ giải thích, phù hợp dữ liệu vừa và nhỏ. Hướng học sâu dùng mạng tích chập (ResNet, EfficientNet, "
        "MobileNet) hoặc Transformer để tự học đặc trưng, thường cho độ chính xác cao hơn khi có dữ liệu lớn và GPU. "
        "Các nghiên cứu gần đây còn kết hợp segmentation để tách vùng thịt và bổ sung cơ chế phát hiện mẫu ngoài phân "
        "phối (OOD) nhằm tăng độ tin cậy khi triển khai thực tế. Báo cáo này đi theo hướng đặc trưng thủ công + "
        "RandomForest như một baseline mạnh, đồng thời áp dụng ý tưởng tách vùng thịt để giảm nhiễu nền.",
    )
    add_table(
        doc,
        ["Tiêu chí", "Đặc trưng thủ công + ML", "Học sâu (CNN/Transformer)"],
        [
            ["Nhu cầu dữ liệu", "Vừa và nhỏ", "Lớn"],
            ["Tài nguyên", "CPU là đủ", "Thường cần GPU"],
            ["Khả năng giải thích", "Cao (đặc trưng rõ ràng)", "Thấp (hộp đen)"],
            ["Độ chính xác tiềm năng", "Khá – cao", "Cao khi đủ dữ liệu"],
            ["Thời gian triển khai", "Nhanh", "Lâu hơn"],
        ],
        [2600, 3380, 3380],
    )
    add_paragraph(
        doc,
        "Bảng trên cho thấy vì sao baseline đặc trưng thủ công + RandomForest là lựa chọn hợp lý cho phạm vi một bài tập "
        "môn học: nhẹ, dễ giải thích và vẫn đạt độ chính xác cao trên bộ dữ liệu đang xét.",
    )
    add_paragraph(
        doc,
        "Một điểm chung của các nghiên cứu là nhấn mạnh vai trò của màu sắc — đặc biệt độ đỏ liên quan tới myoglobin — "
        "như tín hiệu chủ đạo, đồng thời cảnh báo rủi ro khi mô hình học nhầm đặc trưng nền hoặc điều kiện chụp. Đây "
        "chính là lý do báo cáo áp dụng bước tách vùng thịt và dành phần thảo luận cho hiện tượng lệch phân phối.",
    )

    # ------------------------------------------------------------- 4. DU LIEU
    add_heading(doc, "4. Bộ dữ liệu", 1)
    add_heading(doc, "4.1 Giới thiệu LocBeef", 2)
    add_paragraph(
        doc,
        "Bộ dữ liệu sử dụng là LocBeef — Beef Quality Image Dataset (Kaggle, tác giả mexwell), gồm 3.268 ảnh thịt bò "
        "địa phương vùng Aceh, được gán hai nhãn fresh và rotten. Ảnh chụp miếng thịt đặt trên đĩa/nền gỗ trong điều "
        "kiện ánh sáng tương đối đồng nhất. Bộ dữ liệu đã được chia sẵn thành tập train và tập test. Trong báo cáo, "
        "nhãn rotten được ánh xạ sang spoiled để thống nhất thuật ngữ với ứng dụng.",
    )
    add_table(
        doc,
        ["Tập", "fresh", "rotten", "Tổng"],
        [
            ["Train", "1.144", "1.144", "2.288"],
            ["Test", "490", "490", "980"],
            ["Tổng", "1.634", "1.634", "3.268"],
        ],
        [2400, 2320, 2320, 2320],
    )
    add_paragraph(
        doc,
        "Hai lớp cân bằng hoàn hảo (mỗi lớp 1.634 ảnh), nên accuracy là chỉ số đánh giá hợp lý và không bị lệch do mất "
        "cân bằng lớp.",
    )
    add_figure(doc, FIG / "dataset_samples.png",
               "Hình 1. Ảnh mẫu từ LocBeef. Hàng trên: thịt tươi (đỏ tươi). Hàng dưới: thịt hỏng (nâu sẫm, xỉn màu).",
               width=6.2)
    add_figure(doc, FIG / "eda_scatter.png",
               "Hình 2. Phân bố màu vùng thịt trên mặt phẳng (a*, L*): hai lớp tách biệt khá rõ, ủng hộ việc dùng đặc trưng màu.",
               width=5.4)
    add_heading(doc, "4.2 Đặc điểm và thách thức", 2)
    add_bullet(doc, "Nền và vật thể phụ: đĩa trắng, mặt bàn gỗ, dấu mốc thời gian trên ảnh — có thể gây nhiễu nếu không tách vùng thịt.")
    add_bullet(doc, "Khác biệt độ sáng giữa các ảnh do thời điểm chụp, cần cân bằng sáng.")
    add_bullet(doc, "Ranh giới tươi/hỏng mang tính liên tục; một số mẫu ở vùng chuyển tiếp có thể khó phân loại.")
    add_paragraph(
        doc,
        "Do kho ảnh gốc lớn (khoảng 5,5 GB), quá trình huấn luyện đọc ảnh trực tiếp từ file nén trong bộ nhớ mà không "
        "giải nén ra ổ đĩa, giúp tiết kiệm dung lượng (script scripts/train_locbeef_from_zip.py).",
    )

    # ---------------------------------------------------------- 5. PHUONG PHAP
    add_heading(doc, "5. Phương pháp đề xuất", 1)
    add_heading(doc, "5.1 Tổng quan pipeline", 2)
    for step in [
        "Đọc ảnh và resize về 224 × 224 để chuẩn hóa kích thước đầu vào.",
        "Cân bằng sáng cục bộ bằng CLAHE trên kênh L của không gian Lab.",
        "Tách vùng thịt (meat-region masking) để loại nền trắng/sáng và vùng quá tối.",
        "Trích đặc trưng màu (histogram HSV, Lab, thống kê) và texture (LBP) chỉ trên vùng thịt.",
        "Ghép thành vector 174 chiều và phân loại bằng RandomForest.",
        "Đánh giá bằng accuracy, precision, recall, F1-score và confusion matrix.",
    ]:
        add_number(doc, step)
    add_figure(doc, FIG / "pipeline_diagram.png",
               "Hình 3. Sơ đồ khối tổng quan của pipeline nhận biết thịt tươi.", width=6.8)

    add_heading(doc, "5.2 Tiền xử lý và tách vùng thịt", 2)
    add_paragraph(
        doc,
        "Sau khi cân bằng sáng bằng CLAHE, hệ thống tạo một mask nhị phân giữ lại pixel thuộc miếng thịt. Quy tắc: loại "
        "các pixel quá sáng và ít bão hòa (S < 30 và V > 220 — tương ứng đĩa/nền trắng) và pixel quá tối (V ≤ 20 — vùng "
        "bóng). Mask sau đó được làm sạch bằng phép hình thái học mở rồi đóng (open + close) với phần tử cấu trúc elip "
        "9×9 để loại đốm nhiễu và lấp lỗ nhỏ. Nếu vùng thịt phát hiện được nhỏ hơn 3% diện tích ảnh, hệ thống bỏ mask và "
        "dùng toàn ảnh để tránh trường hợp mất dữ liệu. Toàn bộ đặc trưng màu chỉ được tính trên các pixel thuộc vùng "
        "thịt, nhờ đó mô hình không học nhầm màu nền hay màu đĩa.",
    )
    add_paragraph(
        doc,
        "Phép mở (open) giúp loại các đốm nhiễu nhỏ và mảng sáng lẻ tẻ bị nhận nhầm là thịt, trong khi phép đóng (close) "
        "lấp các lỗ nhỏ bên trong miếng thịt (do phản chiếu hay vân mỡ). Thứ tự mở-rồi-đóng cho một mask liền mạch và "
        "sạch hơn, làm đầu vào ổn định cho bước trích đặc trưng. Cơ chế dự phòng dùng toàn ảnh khi mask quá nhỏ đảm bảo "
        "hệ thống không thất bại với các ảnh khó phân đoạn.",
    )
    add_figure(doc, FIG / "preprocessing.png",
               "Hình 4. Quy trình tiền xử lý: ảnh gốc → sau CLAHE → mask vùng thịt → vùng thịt được giữ lại.",
               width=6.6)

    add_heading(doc, "5.3 Trích đặc trưng", 2)
    add_paragraph(
        doc,
        "Vector đặc trưng gồm bốn nhóm, tổng cộng 174 chiều, mô tả trong bảng dưới. Histogram từng kênh được chuẩn hóa "
        "theo mật độ để bất biến với số lượng pixel vùng thịt.",
    )
    add_table(
        doc,
        ["Nhóm đặc trưng", "Chi tiết", "Số chiều"],
        [
            ["Histogram HSV", "H (32 bins, 0–179), S (16 bins), V (16 bins)", "64"],
            ["Histogram Lab", "L (16), a (16), b (16)", "48"],
            ["Thống kê màu", "mean, std, phân vị 10/50/90 trên 6 kênh HSV+Lab", "30"],
            ["Texture LBP", "Local Binary Pattern, 32 bins", "32"],
            ["Tổng cộng", "", "174"],
        ],
        [3000, 5060, 1300],
    )
    add_paragraph(
        doc,
        "Đặc trưng màu chiếm phần lớn vector vì màu là tín hiệu phân biệt mạnh nhất giữa thịt tươi và hỏng. Hình 3 minh "
        "họa: phân bố giá trị trung bình kênh a* (độ đỏ) của thịt tươi lệch rõ về phía cao so với thịt hỏng, xác nhận "
        "cơ sở của việc dùng đặc trưng màu.",
    )
    add_figure(doc, FIG / "color_distribution.png",
               "Hình 5. Phân bố độ đỏ a* (CIELAB) trên vùng thịt: thịt tươi đỏ hơn rõ rệt so với thịt hỏng.",
               width=5.6)

    add_heading(doc, "5.4 Huấn luyện mô hình", 2)
    add_paragraph(
        doc,
        "Mô hình chính là RandomForest với 300 cây, class_weight='balanced', random_state cố định để tái lập. Để đo "
        "hiệu năng một cách khách quan, mô hình được huấn luyện trên tập train (2.288 ảnh) và đánh giá trên tập test "
        "(980 ảnh) vốn không xuất hiện khi huấn luyện. Riêng mô hình đóng gói kèm ứng dụng (để triển khai) được huấn "
        "luyện lại trên toàn bộ 3.268 ảnh nhằm tận dụng tối đa dữ liệu; con số accuracy báo cáo vẫn lấy từ thí nghiệm "
        "train/test có held-out ở trên.",
    )
    add_paragraph(
        doc,
        "Cách làm này tuân theo thực hành phổ biến trong học máy: dùng phần held-out để ước lượng khả năng tổng quát một "
        "cách trung thực, rồi huấn luyện mô hình triển khai trên toàn bộ dữ liệu để đạt hiệu năng tốt nhất có thể. Nhờ "
        "quá trình đọc ảnh trực tiếp từ file nén, việc huấn luyện lại trên toàn bộ dữ liệu không đòi hỏi thêm dung "
        "lượng lưu trữ.",
    )

    # ------------------------------------------------------- 6. KET QUA
    add_heading(doc, "6. Thực nghiệm và kết quả", 1)
    add_heading(doc, "6.1 Thiết lập", 2)
    add_paragraph(
        doc,
        "Thí nghiệm dùng phân chia train/test có sẵn của LocBeef. Ảnh được xử lý qua đúng pipeline mô tả ở Phần 5. Các "
        "chỉ số được tính trên 980 ảnh test.",
    )
    add_heading(doc, "6.2 Kết quả tổng thể", 2)
    add_table(
        doc,
        ["Chỉ số", "fresh", "spoiled"],
        [
            ["Precision", "1.00", "0.96"],
            ["Recall", "0.96", "1.00"],
            ["F1-score", "0.98", "0.98"],
        ],
        [3200, 3080, 3080],
    )
    add_paragraph(
        doc,
        "Accuracy tổng thể đạt 97.9% (959/980 ảnh đúng, 21 lỗi). Cả hai lớp đều đạt F1-score 0.98, cho thấy mô hình cân "
        "bằng tốt giữa tươi và hỏng.",
        bold_prefix="Accuracy tổng thể",
    )
    add_heading(doc, "6.3 Ma trận nhầm lẫn", 2)
    add_paragraph(
        doc,
        "Confusion matrix (Hình 4) cho thấy toàn bộ 490 ảnh hỏng được nhận đúng; 21 ảnh tươi bị đoán nhầm thành hỏng và "
        "không có ảnh hỏng nào bị đoán nhầm thành tươi. Nói cách khác, mô hình nghiêng về phía thận trọng: khi sai, nó "
        "sai theo hướng cảnh báo hỏng chứ không bỏ sót thịt hỏng — đặc tính có lợi cho bài toán an toàn thực phẩm.",
    )
    add_figure(doc, CONFUSION_MATRIX,
               "Hình 6. Confusion matrix trên tập test LocBeef (980 ảnh).", width=4.6)
    add_heading(doc, "6.4 Độ quan trọng đặc trưng", 2)
    add_paragraph(
        doc,
        "Tổng hợp độ quan trọng của RandomForest theo nhóm đặc trưng (Hình 5) cho thấy các đặc trưng màu (histogram "
        "HSV, histogram Lab và thống kê màu) đóng góp phần lớn quyết định, trong khi texture LBP đóng vai trò bổ trợ. "
        "Kết quả này phù hợp với trực giác: màu là dấu hiệu chính để phân biệt thịt tươi và hỏng.",
    )
    add_figure(doc, FIG / "feature_importance.png",
               "Hình 7. Đóng góp của từng nhóm đặc trưng vào mô hình RandomForest.", width=5.6)
    add_heading(doc, "6.5 Phân tích lỗi", 2)
    add_paragraph(
        doc,
        "21 lỗi đều là ảnh tươi bị phân loại thành hỏng. Quan sát cho thấy các mẫu này thường có miếng thịt sẫm màu hơn "
        "trung bình hoặc bị bóng/thiếu sáng cục bộ khiến độ đỏ đo được thấp đi. Đây là các trường hợp ở vùng chuyển "
        "tiếp về màu; hướng khắc phục gồm tăng cường dữ liệu và bổ sung đặc trưng phản ánh độ đồng nhất màu trong vùng "
        "thịt.",
    )
    add_heading(doc, "6.6 Ngưỡng quyết định và điểm vận hành", 2)
    add_paragraph(
        doc,
        "Mô hình xuất xác suất cho mỗi lớp; nhãn được quyết định theo ngưỡng 0,5. Trong thực tế an toàn thực phẩm, có "
        "thể hạ ngưỡng để mô hình 'nhạy' hơn với thịt hỏng (tăng recall lớp hỏng), đổi lại chấp nhận nhiều cảnh báo "
        "nhầm hơn (giảm precision). Việc chọn điểm vận hành nên dựa trên chi phí tương đối giữa bỏ sót thịt hỏng và báo "
        "nhầm thịt tươi trong ứng dụng cụ thể.",
    )
    add_heading(doc, "6.7 Kiểm tra khả năng tổng quát hóa", 2)
    add_paragraph(
        doc,
        "Để đánh giá khả năng tổng quát ngoài miền huấn luyện, mô hình được thử trên một số ảnh thịt lấy ngẫu nhiên từ "
        "web (khác camera, ánh sáng và cách trình bày so với LocBeef). Kết quả cho thấy độ tin cậy giảm rõ và xuất "
        "hiện lỗi — ví dụ một ảnh thịt đã ngả màu bị phân loại là tươi. Điều này khẳng định mô hình mạnh trên ảnh giống "
        "phân phối huấn luyện nhưng chưa tổng quát cho mọi điều kiện; đây là căn cứ cho các hạn chế nêu ở Phần 9.",
    )

    # -------------------------------------------------------- 7. WEB DEMO
    add_heading(doc, "7. Ứng dụng web demo", 1)
    add_heading(doc, "7.1 Kiến trúc", 2)
    add_paragraph(
        doc,
        "Ứng dụng được xây dựng bằng Flask (app.py) với giao diện HTML/JavaScript (templates/index.html). Máy chủ nạp "
        "sẵn mô hình một lần và phục vụ hai endpoint: trang chính để tải ảnh, và endpoint /predict nhận ảnh, chạy pipeline "
        "trích đặc trưng + RandomForest, trả về JSON gồm nhãn, mô tả tiếng Việt, độ tin cậy và xác suất từng lớp. Endpoint "
        "/health dùng để kiểm tra tình trạng mô hình.",
    )
    add_table(
        doc,
        ["Endpoint", "Phương thức", "Chức năng"],
        [
            ["/", "GET", "Trả về trang giao diện tải ảnh."],
            ["/predict", "POST", "Nhận ảnh (multipart), trả JSON nhãn + xác suất."],
            ["/health", "GET", "Kiểm tra mô hình đã nạp và danh sách lớp."],
        ],
        [2200, 2200, 4960],
    )
    add_paragraph(doc, "Ví dụ phản hồi JSON của /predict:")
    code = doc.add_paragraph()
    code_run = code.add_run(
        '{ "label": "spoiled", "label_vi": "Hỏng / Ôi thiu", "level": "bad",\n'
        '  "confidence": 0.94, "probabilities": { "fresh": 0.06, "spoiled": 0.94 } }'
    )
    set_font(code_run, name="Consolas", size=9)
    add_heading(doc, "7.2 Luồng sử dụng", 2)
    for step in [
        "Người dùng mở trình duyệt tại http://127.0.0.1:5000.",
        "Kéo-thả hoặc chọn ảnh miếng thịt; giao diện hiển thị ảnh xem trước.",
        "Nhấn 'Phân tích'; ảnh được gửi tới máy chủ.",
        "Kết quả hiển thị: nhãn (Tươi/Hỏng) với màu tương ứng, độ tin cậy và thanh xác suất từng lớp.",
    ]:
        add_number(doc, step)

    add_heading(doc, "7.3 Minh họa giao diện web", 2)
    add_paragraph(
        doc,
        "Phần này trình bày các ảnh chụp màn hình thực tế của ứng dụng đang chạy trên trình duyệt. Giao diện được thiết "
        "kế tối giản, hai cột: cột trái để tải và xem trước ảnh, cột phải hiển thị kết quả. Bảng màu tối giúp làm nổi "
        "bật ảnh thịt và các thanh xác suất.",
    )
    add_paragraph(
        doc,
        "Khi mới mở, trang hiển thị vùng kéo-thả cùng hướng dẫn định dạng và giới hạn dung lượng (Hình 8). Người dùng "
        "có thể bấm chọn tệp hoặc kéo ảnh trực tiếp vào vùng này.",
    )
    add_figure(doc, FIG / "web_upload.png",
               "Hình 8. Màn hình ban đầu: vùng kéo-thả ảnh và ghi chú định dạng, dung lượng.", width=6.2)
    add_paragraph(
        doc,
        "Sau khi chọn ảnh và nhấn 'Phân tích', ảnh xem trước hiện ở cột trái, còn cột phải hiển thị nhãn kết quả kèm "
        "màu (xanh cho Tươi), độ tin cậy và thanh xác suất cho từng lớp. Hình 9 minh họa một ảnh thịt tươi được dự "
        "đoán đúng là 'Tươi' với độ tin cậy cao; hai thanh xác suất cho thấy phần lớn khối lượng dồn vào lớp Tươi.",
    )
    add_figure(doc, FIG / "web_result_fresh.png",
               "Hình 9. Kết quả cho ảnh thịt tươi: nhãn 'Tươi', độ tin cậy và thanh xác suất từng lớp.", width=6.6)
    add_paragraph(
        doc,
        "Với ảnh thịt đã hỏng, hệ thống trả về nhãn 'Hỏng / Ôi thiu' (màu đỏ) kèm mô tả cảnh báo và thanh xác suất "
        "nghiêng hẳn về lớp Hỏng (Hình 10). Cách phối màu theo mức độ (xanh/vàng/đỏ) giúp người dùng nắm nhanh kết quả "
        "mà không cần đọc kỹ con số.",
    )
    add_figure(doc, FIG / "web_result_spoiled.png",
               "Hình 10. Kết quả cho ảnh thịt hỏng: nhãn 'Hỏng / Ôi thiu' với xác suất cao ở lớp Hỏng.", width=6.6)
    add_paragraph(
        doc,
        "Toàn bộ tương tác diễn ra không tải lại trang: ảnh được gửi bất đồng bộ tới endpoint /predict và kết quả được "
        "dựng lại bằng JavaScript. Phía dưới trang luôn có ghi chú nhắc rằng công cụ chỉ hỗ trợ sàng lọc, không thay "
        "thế kiểm nghiệm chính thức.",
    )

    add_heading(doc, "7.4 Kết quả demo trên tập test", 2)
    add_paragraph(
        doc,
        "Ngoài các ảnh chụp giao diện, Hình 11 tổng hợp dự đoán của mô hình trên sáu ảnh test LocBeef: ba ảnh tươi và "
        "ba ảnh hỏng. Tất cả đều được phân loại đúng với độ tin cậy cao; thịt tươi có màu đỏ tươi, thịt hỏng có màu nâu "
        "sẫm rõ rệt.",
    )
    add_figure(doc, FIG / "demo_predictions.png",
               "Hình 11. Demo dự đoán trên ảnh test: dấu ✓ xanh là dự đoán đúng, kèm độ tin cậy.", width=6.4)
    add_heading(doc, "7.5 Xử lý định dạng ảnh và độ bền", 2)
    add_paragraph(
        doc,
        "Ảnh chụp từ điện thoại hoặc tải trên web đôi khi ở định dạng AVIF/HEIC (thường vẫn mang đuôi .jpg) mà OpenCV "
        "không giải mã được. Ứng dụng bổ sung cơ chế dự phòng: khi OpenCV thất bại, hệ thống chuyển sang Pillow để đọc "
        "ảnh, nhờ đó không báo lỗi với các định dạng hiện đại. Ứng dụng cũng kiểm tra phần mở rộng và giới hạn dung "
        "lượng tải lên (10 MB).",
    )
    add_heading(doc, "7.6 Cách chạy", 2)
    add_paragraph(
        doc,
        "Cài phụ thuộc bằng pip install -r requirements.txt rồi chạy python app.py. Mô hình đã được đóng gói sẵn trong "
        "repo nên không cần huấn luyện lại. Có thể trỏ sang mô hình khác qua biến môi trường MEAT_MODEL.",
    )

    # ------------------------------------------------ 8. CAI DAT & TAI LAP
    add_heading(doc, "8. Chi tiết cài đặt và khả năng tái lập", 1)
    add_heading(doc, "8.1 Công nghệ sử dụng", 2)
    add_paragraph(
        doc,
        "Hệ thống được cài đặt bằng Python. Xử lý ảnh và trích đặc trưng dùng OpenCV và NumPy; huấn luyện và đánh giá "
        "dùng scikit-learn (RandomForest, SVM, LabelEncoder, các chỉ số). Biểu đồ dùng Matplotlib. Ứng dụng web dùng "
        "Flask, ảnh định dạng mới được đọc bằng Pillow. Mô hình được lưu bằng joblib. Toàn bộ chạy trên CPU, không cần GPU.",
    )
    add_heading(doc, "8.2 Cấu trúc mã nguồn", 2)
    add_table(
        doc,
        ["Thành phần", "Vai trò"],
        [
            ["app.py", "Ứng dụng web Flask: tải ảnh, dự đoán, trả JSON."],
            ["templates/index.html", "Giao diện web (kéo-thả, xem trước, thanh xác suất)."],
            ["src/features.py", "Tách vùng thịt và trích đặc trưng 174 chiều."],
            ["src/train.py", "Huấn luyện SVM/RandomForest từ thư mục ảnh."],
            ["src/predict.py", "Dự đoán một ảnh từ dòng lệnh."],
            ["scripts/train_locbeef_from_zip.py", "Train từ file zip LocBeef, không giải nén."],
            ["scripts/eval_locbeef_from_zip.py", "Đánh giá + xuất confusion matrix, report."],
            ["scripts/make_report_figures.py", "Sinh các hình minh họa cho báo cáo."],
            ["scripts/build_report.py", "Sinh file báo cáo Word."],
        ],
        [3600, 5760],
    )
    add_heading(doc, "8.3 Quy trình tái lập kết quả", 2)
    for step in [
        "Đánh giá bằng train/test split để lấy số liệu: python scripts/eval_locbeef_from_zip.py --zip archive.zip.",
        "Huấn luyện mô hình cuối cùng trên toàn bộ dữ liệu: python scripts/train_locbeef_from_zip.py --zip archive.zip --all.",
        "Sinh hình minh họa: python scripts/make_report_figures.py --zip archive.zip.",
        "Sinh báo cáo Word: python scripts/build_report.py.",
        "Chạy ứng dụng web: python app.py rồi mở http://127.0.0.1:5000.",
    ]:
        add_number(doc, step)
    add_paragraph(
        doc,
        "Việc cố định random_state trong RandomForest và dùng đúng phân chia train/test có sẵn giúp kết quả tái lập "
        "ổn định giữa các lần chạy.",
    )

    # ---------------------------------------------------------- 9. THAO LUAN
    add_heading(doc, "9. Thảo luận", 1)
    add_heading(doc, "9.1 Ưu điểm", 2)
    add_bullet(doc, "Pipeline đơn giản, chạy nhanh trên CPU, không cần GPU; phù hợp bài tập và prototype.")
    add_bullet(doc, "Kết quả cao trên đúng phân phối dữ liệu huấn luyện (97.9%).")
    add_bullet(doc, "Có thể giải thích: quyết định dựa trên đặc trưng màu/texture rõ ràng, kèm độ quan trọng đặc trưng.")
    add_heading(doc, "9.2 Hạn chế và điểm cần trung thực", 2)
    add_bullet(
        doc,
        "Cách chia train/test dùng phân chia có sẵn của bộ dữ liệu. Nếu nhiều ảnh chụp cùng một miếng thịt xuất hiện ở "
        "cả train lẫn test, con số 97.9% có thể lạc quan hơn so với thực tế trên mẫu hoàn toàn mới. Đây là giới hạn cần "
        "lưu ý khi diễn giải kết quả.",
        bold_prefix="Rò rỉ dữ liệu tiềm ẩn: ",
    )
    add_bullet(
        doc,
        "Mô hình được huấn luyện trên ảnh bò Aceh chụp khá đồng nhất. Khi thử với ảnh thịt lấy ngẫu nhiên trên web "
        "(khác camera, ánh sáng, loại thịt), độ tin cậy giảm và có trường hợp sai. Mô hình mạnh trên ảnh giống dữ liệu "
        "huấn luyện, không nên xem là bộ phân loại tổng quát cho mọi loại thịt.",
        bold_prefix="Lệch phân phối (domain shift): ",
    )
    add_bullet(
        doc,
        "Trong quá trình phát triển, một lớp hậu xử lý 'hybrid' phân tích màu miền (CIELAB/HSV) từng được thêm vào. "
        "Tuy nhiên khi đánh giá định lượng trên tập test thật, lớp này kéo accuracy xuống 50% do các ngưỡng màu được "
        "chỉnh tay theo ảnh stock sáng đẹp, không khớp màu bò thật nên ép hầu hết ảnh về 'hỏng'. Lớp hybrid đã bị loại "
        "bỏ; hệ thống dùng trực tiếp dự đoán của RandomForest. Đây là minh chứng cho tầm quan trọng của việc kiểm chứng "
        "heuristic bằng số liệu thay vì tin vào cảm tính.",
        bold_prefix="Bài học từ lớp hybrid: ",
    )
    add_bullet(
        doc,
        "Nhãn tươi/hỏng của bộ dữ liệu mang tính nhị phân trong khi độ tươi thực tế biến thiên liên tục; ranh giới gán "
        "nhãn phụ thuộc quy ước của người thu thập dữ liệu, có thể khác nhau giữa các nguồn.",
        bold_prefix="Tính chủ quan của nhãn: ",
    )
    add_heading(doc, "9.3 Khía cạnh an toàn", 2)
    add_paragraph(
        doc,
        "Hệ thống là công cụ hỗ trợ sàng lọc nhanh, không thay thế kiểm nghiệm vi sinh hay đánh giá an toàn thực phẩm "
        "chính thức. Dự đoán 'tươi' không bảo đảm an toàn tuyệt đối; quyết định sử dụng thực phẩm cần dựa trên quy trình "
        "kiểm định phù hợp. Vì lý do này, giao diện web luôn kèm ghi chú nhắc người dùng về giới hạn của công cụ.",
    )

    # ----------------------------------------------------------- 10. KET LUAN
    add_heading(doc, "10. Kết luận và hướng phát triển", 1)
    add_paragraph(
        doc,
        "Báo cáo đã xây dựng một hệ thống thị giác máy tính nhận biết độ tươi thịt bò từ ảnh, gồm pipeline tiền xử lý, "
        "tách vùng thịt, trích đặc trưng màu + texture 174 chiều và phân loại bằng RandomForest. Trên bộ dữ liệu thật "
        "LocBeef, hệ thống đạt 97.9% accuracy trên tập test và được đóng gói thành ứng dụng web tiện dụng.",
    )
    add_paragraph(doc, "Các hướng phát triển tiếp theo:")
    add_bullet(doc, "Chia dữ liệu theo từng mẫu vật để loại trừ hoàn toàn khả năng rò rỉ và ước lượng khả năng tổng quát chính xác hơn.")
    add_bullet(doc, "Tăng cường và đa dạng hóa dữ liệu (nhiều loại thịt, nhiều điều kiện chụp) để giảm lệch phân phối.")
    add_bullet(doc, "Áp dụng transfer learning (MobileNetV3, EfficientNet-B0) và so sánh với baseline hiện tại.")
    add_bullet(doc, "Bổ sung cơ chế từ chối dự đoán khi ảnh nằm ngoài phân phối hoặc độ tin cậy thấp, tăng an toàn khi triển khai.")

    # -------------------------------------------------------- TAI LIEU
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Tài liệu tham khảo", 1)
    refs = [
        "LocBeef — Beef Quality Image Dataset (local Aceh beef, fresh/rotten). Kaggle. https://www.kaggle.com/datasets/mexwell/locbeef-beef-quality-image-dataset",
        "Bramantyo, H. A., Faridi, M. A., Chen, R., Harris, C., & Sun, Y. (2026). Deep Learning-Based Meat Freshness Detection with Segmentation and OOD-Aware Classification. arXiv:2603.00368.",
        "Hidalgo, M. M., Lima, R. C., De Nadai Fernandes, E. A., Bacchi, M. A., & Sarriés, G. A. (2025). Leveraging pre-trained computer vision models for accurate classification of meat freshness. Food Chemistry, 495(Pt 3), 146430.",
        "Ojala, T., Pietikäinen, M., & Mäenpää, T. (2002). Multiresolution gray-scale and rotation invariant texture classification with local binary patterns. IEEE TPAMI, 24(7), 971–987.",
        "Zuiderveld, K. (1994). Contrast Limited Adaptive Histogram Equalization. Graphics Gems IV, 474–485.",
        "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
    ]
    for ref in refs:
        add_number(doc, ref)

    add_heading(doc, "Phụ lục A. Siêu tham số và cấu hình", 1)
    add_table(
        doc,
        ["Thành phần", "Giá trị"],
        [
            ["Kích thước ảnh", "224 × 224"],
            ["CLAHE", "clipLimit = 2.0, tileGridSize = 8 × 8 (trên kênh L)"],
            ["Số chiều đặc trưng", "174 (HSV 64 + Lab 48 + thống kê 30 + LBP 32)"],
            ["Mô hình", "RandomForest, 300 cây, class_weight = balanced"],
            ["random_state", "42 (cố định để tái lập)"],
            ["Chia dữ liệu", "Train 2.288 / Test 980 (theo bộ dữ liệu)"],
            ["Định dạng ảnh hỗ trợ", "JPG, PNG, WEBP, BMP; AVIF/HEIC qua Pillow"],
            ["Giới hạn tải lên (web)", "10 MB"],
        ],
        [3600, 5760],
    )
    add_paragraph(
        doc,
        "Môi trường: Python 3.12, OpenCV, NumPy, scikit-learn, Matplotlib, Flask, Pillow, joblib. Toàn bộ pipeline chạy "
        "trên CPU. Danh sách phụ thuộc đầy đủ ở tệp requirements.txt.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Báo cáo bài tập Thị giác máy tính — Nhận biết thịt tươi")
    set_font(footer_run, size=9)
    footer_run.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build()
